"""
Reportify Pro v2.1 - Enterprise Report Generator
================================================

Complete professional report generation system with 21 templates
across 8 IT job role domains.

Author: Portfolio Showcase
Version: 2.1.0
"""

import sys
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field
from typing import Dict, List
from enum import Enum

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx required. Install: pip install python-docx")
    sys.exit(1)

# ════════════════════════════════════════════════════════════════════════
# ENHANCED TEMPLATE DATABASE (21 Templates)
# ════════════════════════════════════════════════════════════════════════


class ReportCategory(str, Enum):
    SECURITY = "Security"
    DEVOPS = "DevOps"
    CLOUD = "Cloud"
    SYSADMIN = "System Admin"
    PROJECT = "Project Management"
    COMPLIANCE = "Compliance"
    NETWORK = "Networking"
    DATA_ANALYTICS = "Data Analytics"


@dataclass
class ReportData:
    template_key: str = ""
    category: str = ""
    title: str = ""
    subtitle: str = ""
    author: str = ""
    date: str = datetime.now().strftime("%B %d, %Y")
    company_name: str = "Your Company Inc."

    # Core sections
    executive_summary: str = ""
    background: str = ""
    objectives: List[str] = field(default_factory=list)
    scope: str = ""
    methodology: str = ""
    findings: List[str] = field(default_factory=list)
    analysis: str = ""
    recommendations: List[str] = field(default_factory=list)
    conclusion: str = ""
    next_steps: List[str] = field(default_factory=list)

    # Additional fields for specialized templates
    kpis: Dict[str, str] = field(default_factory=dict)
    timeline_entries: List[str] = field(default_factory=list)
    risks: List[str] = field(default_factory=list)
    metadata: Dict[str, str] = field(default_factory=dict)
    appendix: str = ""

    tags: List[str] = field(default_factory=list)


REPORT_TEMPLATES = {
    # ═══ SECURITY DOMAIN (6 templates) ═══
    "vulnerability_assessment": {
        "name": "Vulnerability Assessment Report",
        "category": ReportCategory.SECURITY,
        "icon": "🔒",
        "sections": ["executive_summary", "scope", "methodology", "findings", "risks", "recommendations"],
        "description": "Comprehensive security vulnerability assessment with CVSS ratings",
        "standards": "NIST CSF, CIS Benchmarks"
    },

    "penetration_test": {
        "name": "Penetration Testing Report",
        "category": ReportCategory.SECURITY,
        "icon": "🛡️",
        "sections": ["executive_summary", "scope", "methodology", "findings", "risks", "recommendations"],
        "description": "Full penetration test following OWASP/PTES methodologies",
        "standards": "OWASP Top 10, PTES"
    },

    "iam_audit": {
        "name": "IAM Security Audit",
        "category": ReportCategory.SECURITY,
        "icon": "👤",
        "sections": ["executive_summary", "scope", "findings", "risks", "recommendations"],
        "description": "Identity and Access Management security assessment",
        "standards": "NIST SP 800-63"
    },

    "attack_surface": {
        "name": "Attack Surface Analysis",
        "category": ReportCategory.SECURITY,
        "icon": "🎯",
        "sections": ["scope", "methodology", "findings", "risks", "recommendations", "appendix"],
        "description": "External attack surface enumeration and risk assessment",
        "standards": "NIST CSF Identify/Protect",
        "defaults": {
            "title": "Attack Surface Analysis Report",
            "objectives": [
                "Enumerate all external-facing assets (domains, IPs, cloud resources)",
                "Identify exposed services, open ports, and misconfigurations",
                "Assess risk based on exploitability and business impact",
                "Provide prioritized remediation recommendations"
            ]
        }
    },

    "incident_response": {
        "name": "Security Incident Report",
        "category": ReportCategory.SECURITY,
        "icon": "🚨",
        "sections": ["metadata", "executive_summary", "timeline", "findings", "analysis", "recommendations"],
        "description": "Post-incident review following NIST SP 800-61",
        "standards": "NIST SP 800-61 r2, ISO 27001 A.16",
        "defaults": {
            "title": "Security Incident Report - {{incident_id}}",
            "objectives": [
                "Document complete incident timeline from detection to recovery",
                "Perform root cause analysis using 5 Whys methodology",
                "Assess business impact (users affected, downtime, data exposure)",
                "Identify control gaps and implement preventive measures"
            ]
        }
    },

    "log_analysis": {
        "name": "Log Analysis Report",
        "category": ReportCategory.SECURITY,
        "icon": "📊",
        "sections": ["scope", "methodology", "findings", "recommendations", "appendix"],
        "description": "Security log analysis and anomaly detection",
        "standards": "NIST Logging Practices",
        "defaults": {
            "title": "Log Analysis Report - {{timeframe}}",
            "methodology": "Analyzed system logs using SIEM correlation rules, statistical anomaly detection, and baseline deviation analysis.",
            "objectives": [
                "Identify suspicious authentication patterns and failed login attempts",
                "Detect anomalous network traffic and data exfiltration indicators",
                "Correlate security events across multiple log sources",
                "Tune SIEM rules and reduce false positive rates"
            ]
        }
    },

    # ═══ DEVOPS DOMAIN (4 templates) ═══
    "infrastructure_design": {
        "name": "Infrastructure Design Document",
        "category": ReportCategory.DEVOPS,
        "icon": "🏗️",
        "sections": ["executive_summary", "objectives", "scope", "methodology", "recommendations"],
        "description": "Infrastructure architecture specification and design"
    },

    "deployment_report": {
        "name": "Production Deployment Report",
        "category": ReportCategory.DEVOPS,
        "icon": "🚀",
        "sections": ["executive_summary", "scope", "timeline", "findings", "recommendations"],
        "description": "Production deployment summary with rollback procedures"
    },

    "proof_of_concept": {
        "name": "Proof of Concept (PoC) Report",
        "category": ReportCategory.DEVOPS,
        "icon": "🧪",
        "sections": ["objectives", "scope", "methodology", "findings", "analysis", "recommendations"],
        "description": "Technical PoC validation with success criteria",
        "standards": "---",
        "defaults": {
            "title": "Proof of Concept Report - {{technology}}",
            "objectives": [
                "Validate technical feasibility against defined success criteria",
                "Measure performance, scalability, and integration capabilities",
                "Identify implementation risks and resource requirements",
                "Provide go/no-go recommendation with supporting evidence"
            ]
        }
    },

    "qa_test_report": {
        "name": "QA Test Report",
        "category": ReportCategory.DEVOPS,
        "icon": "✅",
        "sections": ["overview", "objectives", "findings", "analysis", "recommendations", "appendix"],
        "description": "Functional, regression, and performance testing results",
        "standards": "IEEE 829",
        "defaults": {
            "title": "QA Test Report - {{build_version}}",
            "objectives": [
                "Execute comprehensive test suite (functional, regression, integration)",
                "Validate acceptance criteria and user stories completion",
                "Identify defects by severity and provide root cause analysis",
                "Assess release readiness with go/no-go recommendation"
            ]
        }
    },

    # ═══ CLOUD DOMAIN (2 templates) ═══
    "cloud_migration": {
        "name": "Cloud Migration Assessment",
        "category": ReportCategory.CLOUD,
        "icon": "☁️",
        "sections": ["executive_summary", "objectives", "methodology", "findings", "risks", "recommendations"],
        "description": "Cloud migration strategy and execution plan"
    },

    "cost_optimization": {
        "name": "Cloud Cost Optimization",
        "category": ReportCategory.CLOUD,
        "icon": "💰",
        "sections": ["executive_summary", "findings", "analysis", "recommendations"],
        "description": "Cloud infrastructure cost analysis and optimization"
    },

    # ═══ SYSTEM ADMIN DOMAIN (2 templates) ═══
    "system_audit": {
        "name": "System Administration Audit",
        "category": ReportCategory.SYSADMIN,
        "icon": "🖥️",
        "sections": ["scope", "methodology", "findings", "recommendations"],
        "description": "Server and system configuration audit"
    },

    "technology_assessment": {
        "name": "Technology Assessment",
        "category": ReportCategory.SYSADMIN,
        "icon": "🔧",
        "sections": ["background", "scope", "findings", "risks", "recommendations", "appendix"],
        "description": "Technology stack evaluation and modernization roadmap",
        "standards": "---",
        "defaults": {
            "title": "Technology Assessment - {{system_name}}",
            "objectives": [
                "Inventory current technology stack (versions, support status, licensing)",
                "Benchmark against industry best practices and modern patterns",
                "Identify operational and security risks from technical debt",
                "Provide target state architecture with migration roadmap and costs"
            ]
        }
    },

    # ═══ PROJECT MANAGEMENT DOMAIN (4 templates) ═══
    "project_status": {
        "name": "Project Status Report",
        "category": ReportCategory.PROJECT,
        "icon": "📋",
        "sections": ["executive_summary", "timeline", "risks", "findings", "recommendations"],
        "description": "Current project status and progress tracking"
    },

    "project_proposal": {
        "name": "Project Proposal (Executive-Ready)",
        "category": ReportCategory.PROJECT,
        "icon": "📝",
        "sections": ["executive_summary", "background", "scope", "methodology", "recommendations", "conclusion"],
        "description": "Comprehensive project proposal with ROI analysis",
        "standards": "---",
        "defaults": {
            "title": "Project Proposal - {{project_name}}",
            "objectives": [
                "Define problem statement, business impact, and strategic alignment",
                "Present proposed solution with architecture and alternatives comparison",
                "Provide detailed scope, deliverables, and acceptance criteria",
                "Justify investment with ROI analysis and risk mitigation plan"
            ]
        }
    },

    "project_closure": {
        "name": "Project Outcome/Closure Report",
        "category": ReportCategory.PROJECT,
        "icon": "🏁",
        "sections": ["executive_summary", "objectives", "findings", "analysis", "recommendations"],
        "description": "Project completion summary and lessons learned",
        "standards": "---",
        "defaults": {
            "title": "Project Closure Report - {{project_name}}",
            "objectives": [
                "Compare planned objectives vs actual outcomes with KPI deltas",
                "Analyze budget and schedule variances with root causes",
                "Document benefits realization and metrics achieved to date",
                "Capture lessons learned and best practices for future projects"
            ]
        }
    },

    "operational_report": {
        "name": "Operational Report (Monthly)",
        "category": ReportCategory.PROJECT,
        "icon": "📈",
        "sections": ["executive_summary", "kpis", "findings", "recommendations", "next_steps"],
        "description": "Monthly operational metrics and KPI tracking",
        "standards": "SRE KPIs",
        "defaults": {
            "title": "Operational Report - {{month}} {{year}}",
            "objectives": [
                "Report on key operational KPIs (SLO/SLA, MTTR, MTBF, capacity utilization)",
                "Track project milestone burndown and identify blockers",
                "Highlight automation improvements and toil reduction initiatives",
                "Define next period action items with owners and dependencies"
            ]
        }
    },

    # ═══ COMPLIANCE DOMAIN (1 template) ═══
    "compliance_audit": {
        "name": "Compliance Audit Report",
        "category": ReportCategory.COMPLIANCE,
        "icon": "✅",
        "sections": ["executive_summary", "scope", "methodology", "findings", "recommendations"],
        "description": "Regulatory compliance assessment and gap analysis"
    },

    # ═══ NETWORKING DOMAIN (1 template) ═══
    "network_assessment": {
        "name": "Network Security Assessment",
        "category": ReportCategory.NETWORK,
        "icon": "🌐",
        "sections": ["scope", "methodology", "findings", "risks", "recommendations"],
        "description": "Network infrastructure security and performance review"
    },

    # ═══ DATA ANALYTICS DOMAIN (1 template) ═══
    "analytical_report": {
        "name": "Analytical Report (Data Insights)",
        "category": ReportCategory.DATA_ANALYTICS,
        "icon": "📊",
        "sections": ["executive_summary", "methodology", "kpis", "findings", "recommendations"],
        "description": "Business intelligence and data-driven insights report",
        "standards": "---",
        "defaults": {
            "title": "Data Analytics Report - {{analysis_focus}}",
            "methodology": "Performed comprehensive data analysis using statistical methods, trend analysis, and predictive modeling. Data sources validated for accuracy and completeness.",
            "objectives": [
                "Analyze historical data to identify trends, patterns, and anomalies",
                "Define and track key performance indicators (KPIs) with targets",
                "Provide data-driven insights to support business decision-making",
                "Recommend actionable improvements based on quantitative evidence"
            ]
        }
    }
}


# ════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATOR
# ════════════════════════════════════════════════════════════════════════


class DocumentGenerator:
    """Enhanced document generator with all 21 templates"""

    @staticmethod
    def generate_report(data: ReportData, template: dict, output_path: Path):
        doc = Document()
        DocumentGenerator._setup_styles(doc)

        # Cover page
        DocumentGenerator._add_cover_page(doc, data, template)
        doc.add_page_break()

        # Standards/Frameworks (if applicable)
        if template.get('standards') and template['standards'] != '---':
            doc.add_heading("Standards & Frameworks", 2)
            doc.add_paragraph(f"This report follows: {template['standards']}")
            doc.add_paragraph()

        # Executive Summary
        if data.executive_summary:
            doc.add_heading("Executive Summary", 1)
            doc.add_paragraph(data.executive_summary)
            doc.add_paragraph()

        # Objectives
        if data.objectives:
            doc.add_heading("Objectives", 1)
            for obj in data.objectives:
                doc.add_paragraph(obj, style='List Bullet')
            doc.add_paragraph()

        # Scope
        if data.scope:
            doc.add_heading("Scope", 1)
            doc.add_paragraph(data.scope)
            doc.add_paragraph()

        # Methodology
        if data.methodology:
            doc.add_heading("Methodology", 1)
            doc.add_paragraph(data.methodology)
            doc.add_paragraph()

        # KPIs (for applicable reports)
        if data.kpis:
            doc.add_heading("Key Performance Indicators", 1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            table.rows[0].cells[0].text = "KPI"
            table.rows[0].cells[1].text = "Value"
            for kpi_name, kpi_value in data.kpis.items():
                row = table.add_row().cells
                row[0].text = kpi_name
                row[1].text = kpi_value
            doc.add_paragraph()

        # Findings
        if data.findings:
            doc.add_heading("Key Findings", 1)
            for finding in data.findings:
                doc.add_paragraph(finding, style='List Bullet')
            doc.add_paragraph()

        # Analysis
        if data.analysis:
            doc.add_heading("Analysis", 1)
            doc.add_paragraph(data.analysis)
            doc.add_paragraph()

        # Risks
        if data.risks:
            doc.add_heading("Risk Assessment", 1)
            for risk in data.risks:
                doc.add_paragraph(risk, style='List Bullet')
            doc.add_paragraph()

        # Recommendations
        if data.recommendations:
            doc.add_heading("Recommendations", 1)
            for i, rec in enumerate(data.recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style='List Number')
            doc.add_paragraph()

        # Conclusion
        if data.conclusion:
            doc.add_heading("Conclusion", 1)
            doc.add_paragraph(data.conclusion)
            doc.add_paragraph()

        # Next Steps
        if data.next_steps:
            doc.add_heading("Next Steps", 1)
            for step in data.next_steps:
                doc.add_paragraph(step, style='List Bullet')
            doc.add_paragraph()

        # Appendix
        if data.appendix:
            doc.add_page_break()
            doc.add_heading("Appendix", 1)
            doc.add_paragraph(data.appendix)

        # Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    @staticmethod
    def _setup_styles(doc: Document):
        """Professional document styling"""

        # Configure heading styles
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            style.font.name = 'Calibri'
            style.font.color.rgb = RGBColor(31, 78, 120)

        # Normal text
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

    @staticmethod
    def _add_cover_page(doc: Document, data: ReportData, template: dict):
        """Professional cover page with metadata"""
        # Company header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = header.add_run(f"\n\n{data.company_name}\n")
        company_run.font.size = Pt(18)
        company_run.font.bold = True
        company_run.font.color.rgb = RGBColor(31, 78, 120)

        doc.add_paragraph("\n" * 3)

        # Report title
        title = doc.add_paragraph()
        title_run = title.add_run(data.title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 120)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if data.subtitle:
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(data.subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("\n" * 5)

        # Metadata
        metadata = doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata.add_run(f"Author: {data.author}\n")
        metadata.add_run(f"Date: {data.date}\n")
        metadata.add_run(f"Template: {template['name']}\n")
        if template.get('standards'):
            metadata.add_run(f"Standards: {template['standards']}\n")


# CLI Interface

def main():
    print("Reportify Pro v2.1 - 21 Professional Templates")
    print("=" * 60)
    print("\nAvailable templates:")

    for key, template in REPORT_TEMPLATES.items():
        print(f"  {template['icon']} {key:25} - {template['name']}")

    print("\nUsage:")
    print("  python reportify_pro.py <template_key>")
    print("  Example: python reportify_pro.py vulnerability_assessment")


if __name__ == '__main__':
    main()
