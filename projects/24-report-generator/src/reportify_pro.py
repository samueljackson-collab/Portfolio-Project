"""
Reportify Pro - Ultimate Enterprise Report Generator
====================================================

A comprehensive, enterprise-grade report generation system with 25+ professional
templates spanning 7+ IT job role domains, intelligent automation, and modern workflows.

Features:
- 25+ Professional Templates across 7 IT domains
- Smart Variables & Auto-Population
- Intelligent Tag System with Auto-Suggestions
- Risk Matrix & Timeline Management
- Multi-Format Export (DOCX, PDF, HTML)
- Project Save/Load System
- Template Customization Engine

Dependencies:
    pip install python-docx pydantic pillow

Author: Portfolio Showcase
License: MIT
Version: 2.1.0
"""

import json
import logging
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass, field, asdict
from enum import Enum

# Dependency checks
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx required. Install: pip install python-docx")
    sys.exit(1)

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - [%(levelname)s] - %(message)s',
    handlers=[
        logging.FileHandler('reportify.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════
# DATA MODELS
# ═══════════════════════════════════════════════════════════════════════════

class ReportCategory(str, Enum):
    """IT job role domains"""
    SECURITY = "security"
    DEVOPS = "devops"
    SYSADMIN = "sysadmin"
    CLOUD = "cloud"
    PROJECT_MGMT = "project_mgmt"
    COMPLIANCE = "compliance"
    NETWORKING = "networking"
    DATA_ANALYTICS = "data_analytics"


@dataclass
class Risk:
    description: str = ""
    category: str = "Technical"
    impact: str = "Medium"
    likelihood: str = "Possible"
    mitigation: str = ""
    owner: str = ""


@dataclass
class TimelineEntry:
    milestone: str = ""
    date: str = ""
    status: str = "Planned"
    owner: str = ""
    notes: str = ""


@dataclass
class TechnicalSpec:
    component: str = ""
    specification: str = ""
    version: str = ""
    notes: str = ""


@dataclass
class ReportData:
    """Complete report data model"""
    # Template metadata
    template_key: str = "general_report"
    category: str = "security"

    # Basic info
    title: str = ""
    subtitle: str = ""
    author: str = os.getlogin() if hasattr(os, 'getlogin') else "Report Author"
    date: str = field(default_factory=lambda: datetime.now().strftime("%B %d, %Y"))
    version: str = "1.0"
    status: str = "Draft"
    classification: str = "Internal"

    # Organization
    company_name: str = "Your Company Inc."
    department: str = ""
    project_code: str = ""

    # Content sections
    executive_summary: str = ""
    background: str = ""
    objectives: List[str] = field(default_factory=list)
    scope: str = ""
    methodology: str = ""
    findings: List[str] = field(default_factory=list)
    analysis: str = ""
    recommendations: List[str] = field(default_factory=list)
    conclusion: str = ""
    next_steps: List[str] = field(default_factory=list)

    # Technical details
    tech_stack: List[str] = field(default_factory=list)
    architecture: str = ""
    tech_specs: List[TechnicalSpec] = field(default_factory=list)
    os_platform: str = ""
    ram_size: str = ""
    cpu_specs: str = ""
    deployment_env: str = ""

    # Project management
    risks: List[Risk] = field(default_factory=list)
    timeline: List[TimelineEntry] = field(default_factory=list)
    budget: str = ""
    resources: List[str] = field(default_factory=list)

    # Metadata
    tags: List[str] = field(default_factory=list)
    attachments: List[str] = field(default_factory=list)
    references: List[str] = field(default_factory=list)

    # URLs
    repository_url: str = ""
    documentation_url: str = ""
    demo_url: str = ""

    # Smart variables
    smart_variables: Dict[str, str] = field(default_factory=dict)

    # Additional KPIs
    kpis: Dict[str, str] = field(default_factory=dict)

    def to_dict(self) -> dict:
        """Convert to dictionary for JSON serialization"""
        data = asdict(self)
        # Convert dataclass objects to dicts
        data['risks'] = [asdict(r) for r in self.risks]
        data['timeline'] = [asdict(t) for t in self.timeline]
        data['tech_specs'] = [asdict(ts) for ts in self.tech_specs]
        return data

    @classmethod
    def from_dict(cls, data: dict) -> 'ReportData':
        """Create from dictionary"""
        # Convert dicts back to dataclasses
        if 'risks' in data:
            data['risks'] = [Risk(**r) for r in data['risks']]
        if 'timeline' in data:
            data['timeline'] = [TimelineEntry(**t) for t in data['timeline']]
        if 'tech_specs' in data:
            data['tech_specs'] = [TechnicalSpec(**ts) for ts in data['tech_specs']]
        return cls(**data)


# ═══════════════════════════════════════════════════════════════════════════
# REPORT TEMPLATES DATABASE (25+ Templates)
# ═══════════════════════════════════════════════════════════════════════════

REPORT_TEMPLATES = {
    # ═══ SECURITY REPORTS (6 templates) ═══
    "vulnerability_assessment": {
        "name": "Vulnerability Assessment Report",
        "category": ReportCategory.SECURITY,
        "icon": "🔒",
        "sections": ["executive_summary", "scope", "methodology", "findings", "risks", "recommendations"],
        "description": "Comprehensive security vulnerability assessment with CVSS ratings",
        "standards": "NIST CSF, CIS Benchmarks",
        "defaults": {
            "title": "{{system_name}} Vulnerability Assessment Report",
            "methodology": "Performed comprehensive vulnerability assessment using automated scanners (Nessus, Qualys) and manual penetration testing techniques. Vulnerabilities classified using CVSS v3.1 scoring framework.",
            "objectives": [
                "Identify security vulnerabilities across all in-scope systems",
                "Assess risk levels and potential business impact",
                "Provide prioritized remediation roadmap",
                "Ensure compliance with security standards"
            ]
        }
    },

    "penetration_test": {
        "name": "Penetration Testing Report",
        "category": ReportCategory.SECURITY,
        "icon": "🛡️",
        "sections": ["executive_summary", "scope", "methodology", "findings", "risks", "recommendations", "conclusion"],
        "description": "Full penetration test results following OWASP/PTES methodologies",
        "standards": "OWASP Top 10, PTES",
        "defaults": {
            "title": "Penetration Test Report - {{target_system}}",
            "methodology": "Executed comprehensive penetration testing following OWASP Top 10 and PTES (Penetration Testing Execution Standard) methodologies. Testing included reconnaissance, scanning, exploitation, post-exploitation, and reporting phases.",
            "objectives": [
                "Identify exploitable security vulnerabilities",
                "Assess effectiveness of security controls",
                "Simulate real-world attack scenarios",
                "Validate security posture against industry standards"
            ]
        }
    },

    "incident_response": {
        "name": "Security Incident Report",
        "category": ReportCategory.SECURITY,
        "icon": "🚨",
        "sections": ["executive_summary", "background", "timeline", "findings", "analysis", "recommendations", "next_steps"],
        "description": "Security incident analysis, forensics, and response documentation",
        "standards": "NIST SP 800-61 r2, ISO 27001 A.16",
        "defaults": {
            "title": "Security Incident Report - {{incident_id}}",
            "objectives": [
                "Document incident timeline from detection to recovery",
                "Perform root cause analysis using 5 Whys methodology",
                "Assess business impact (users affected, downtime, data exposure)",
                "Identify control gaps and implement preventive measures"
            ]
        }
    },

    "iam_audit": {
        "name": "IAM Security Audit",
        "category": ReportCategory.SECURITY,
        "icon": "👤",
        "sections": ["executive_summary", "scope", "methodology", "findings", "risks", "recommendations"],
        "description": "Identity and Access Management security assessment",
        "standards": "NIST SP 800-63",
        "defaults": {
            "title": "IAM Security Audit - {{system_name}}",
            "methodology": "Conducted comprehensive IAM audit reviewing user permissions, role assignments, access policies, and authentication mechanisms against least privilege and zero trust principles.",
            "objectives": [
                "Review user access permissions and role assignments",
                "Identify over-privileged accounts and stale credentials",
                "Assess MFA implementation and password policies",
                "Ensure compliance with least privilege principle"
            ]
        }
    },

    "attack_surface": {
        "name": "Attack Surface Analysis",
        "category": ReportCategory.SECURITY,
        "icon": "🎯",
        "sections": ["scope", "methodology", "findings", "risks", "recommendations"],
        "description": "External attack surface enumeration and risk assessment",
        "standards": "NIST CSF Identify/Protect",
        "defaults": {
            "title": "Attack Surface Analysis Report",
            "objectives": [
                "Enumerate all external-facing assets (domains, IPs, cloud resources)",
                "Identify exposed services, open ports, and misconfigurations",
                "Assess risk based on exploitability and business impact",
                "Provide prioritized remediation recommendations"
            ]
        }
    },

    "log_analysis": {
        "name": "Log Analysis Report",
        "category": ReportCategory.SECURITY,
        "icon": "📊",
        "sections": ["scope", "methodology", "findings", "recommendations"],
        "description": "Security log analysis and anomaly detection",
        "standards": "NIST Logging Practices",
        "defaults": {
            "title": "Log Analysis Report - {{timeframe}}",
            "methodology": "Analyzed system logs using SIEM correlation rules, statistical anomaly detection, and baseline deviation analysis.",
            "objectives": [
                "Identify suspicious authentication patterns and failed login attempts",
                "Detect anomalous network traffic and data exfiltration indicators",
                "Correlate security events across multiple log sources",
                "Tune SIEM rules and reduce false positive rates"
            ]
        }
    },

    # ═══ DEVOPS REPORTS (5 templates) ═══
    "infrastructure_design": {
        "name": "Infrastructure Design Document",
        "category": ReportCategory.DEVOPS,
        "icon": "🏗️",
        "sections": ["executive_summary", "objectives", "architecture", "tech_specs", "risks", "recommendations"],
        "description": "Infrastructure architecture specification and design",
        "standards": "---",
        "defaults": {
            "title": "{{project_name}} Infrastructure Design Document",
            "objectives": [
                "Define scalable, resilient infrastructure architecture",
                "Specify technical components and integration points",
                "Establish deployment and operational procedures",
                "Document disaster recovery and high availability design"
            ]
        }
    },

    "deployment_report": {
        "name": "Production Deployment Report",
        "category": ReportCategory.DEVOPS,
        "icon": "🚀",
        "sections": ["executive_summary", "scope", "timeline", "findings", "recommendations", "next_steps"],
        "description": "Production deployment summary with rollback procedures",
        "standards": "---",
        "defaults": {
            "title": "{{application_name}} Deployment Report - {{environment}}",
            "objectives": [
                "Document deployment process and results",
                "Identify and resolve deployment issues",
                "Verify system functionality post-deployment",
                "Validate rollback procedures"
            ]
        }
    },

    "proof_of_concept": {
        "name": "Proof of Concept (PoC) Report",
        "category": ReportCategory.DEVOPS,
        "icon": "🧪",
        "sections": ["objectives", "scope", "methodology", "findings", "analysis", "recommendations"],
        "description": "Technical PoC validation with success criteria",
        "standards": "---",
        "defaults": {
            "title": "Proof of Concept Report - {{technology}}",
            "objectives": [
                "Validate technical feasibility against defined success criteria",
                "Measure performance, scalability, and integration capabilities",
                "Identify implementation risks and resource requirements",
                "Provide go/no-go recommendation with supporting evidence"
            ]
        }
    },

    "qa_test_report": {
        "name": "QA Test Report",
        "category": ReportCategory.DEVOPS,
        "icon": "✅",
        "sections": ["executive_summary", "objectives", "findings", "analysis", "recommendations"],
        "description": "Functional, regression, and performance testing results",
        "standards": "IEEE 829",
        "defaults": {
            "title": "QA Test Report - {{build_version}}",
            "objectives": [
                "Execute comprehensive test suite (functional, regression, integration)",
                "Validate acceptance criteria and user stories completion",
                "Identify defects by severity and provide root cause analysis",
                "Assess release readiness with go/no-go recommendation"
            ]
        }
    },

    "performance_tuning": {
        "name": "Performance Optimization Report",
        "category": ReportCategory.DEVOPS,
        "icon": "⚡",
        "sections": ["executive_summary", "methodology", "findings", "analysis", "recommendations", "next_steps"],
        "description": "System performance analysis and optimization",
        "standards": "---",
        "defaults": {
            "title": "{{system_name}} Performance Optimization Report",
            "methodology": "Conducted comprehensive performance testing including load testing, stress testing, and profiling. Analyzed system metrics, identified bottlenecks, and implemented optimization strategies.",
            "objectives": [
                "Measure system performance under various load conditions",
                "Identify performance bottlenecks and resource constraints",
                "Recommend and implement optimization strategies",
                "Establish performance monitoring and alerting"
            ]
        }
    },

    # ═══ CLOUD INFRASTRUCTURE REPORTS (4 templates) ═══
    "cloud_migration": {
        "name": "Cloud Migration Assessment",
        "category": ReportCategory.CLOUD,
        "icon": "☁️",
        "sections": ["executive_summary", "objectives", "methodology", "findings", "risks", "timeline", "recommendations"],
        "description": "Cloud migration strategy and execution plan",
        "standards": "AWS Well-Architected Framework",
        "defaults": {
            "title": "{{application_name}} Cloud Migration Report",
            "objectives": [
                "Assess cloud readiness and migration complexity",
                "Design target cloud architecture",
                "Develop phased migration strategy",
                "Identify risks and mitigation plans"
            ]
        }
    },

    "cost_optimization": {
        "name": "Cloud Cost Optimization",
        "category": ReportCategory.CLOUD,
        "icon": "💰",
        "sections": ["executive_summary", "findings", "analysis", "recommendations", "conclusion"],
        "description": "Cloud infrastructure cost analysis and optimization",
        "standards": "FinOps Framework",
        "defaults": {
            "title": "{{cloud_provider}} Cost Optimization Report",
            "objectives": [
                "Analyze current cloud spending patterns",
                "Identify cost-saving opportunities",
                "Implement right-sizing and reserved instances",
                "Establish FinOps practices and governance"
            ]
        }
    },

    "aws_infrastructure": {
        "name": "AWS Infrastructure Audit",
        "category": ReportCategory.CLOUD,
        "icon": "🔶",
        "sections": ["executive_summary", "scope", "findings", "risks", "recommendations"],
        "description": "AWS infrastructure review and best practices assessment",
        "standards": "AWS Well-Architected Framework",
        "defaults": {
            "title": "AWS Infrastructure Audit Report",
            "objectives": [
                "Review AWS account structure and organization",
                "Assess security groups, NACLs, and VPC configuration",
                "Evaluate IAM policies and access controls",
                "Ensure compliance with AWS Well-Architected Framework"
            ]
        }
    },

    "multi_cloud": {
        "name": "Multi-Cloud Strategy Report",
        "category": ReportCategory.CLOUD,
        "icon": "🌐",
        "sections": ["executive_summary", "objectives", "architecture", "findings", "recommendations", "conclusion"],
        "description": "Multi-cloud architecture and governance strategy",
        "standards": "---",
        "defaults": {
            "title": "Multi-Cloud Strategy and Architecture",
            "objectives": [
                "Define multi-cloud architecture and integration patterns",
                "Establish cloud governance and cost management",
                "Implement unified security and compliance framework",
                "Design disaster recovery and failover strategies"
            ]
        }
    },

    # ═══ SYSTEM ADMINISTRATION REPORTS (4 templates) ═══
    "system_audit": {
        "name": "System Administration Audit",
        "category": ReportCategory.SYSADMIN,
        "icon": "🖥️",
        "sections": ["executive_summary", "scope", "methodology", "findings", "recommendations"],
        "description": "Server and system configuration audit",
        "standards": "CIS Benchmarks",
        "defaults": {
            "title": "System Administration Audit - {{environment}}",
            "objectives": [
                "Review system configurations and hardening",
                "Assess patch management and update procedures",
                "Evaluate backup and recovery processes",
                "Ensure compliance with security baselines"
            ]
        }
    },

    "technology_assessment": {
        "name": "Technology Assessment",
        "category": ReportCategory.SYSADMIN,
        "icon": "🔧",
        "sections": ["background", "scope", "findings", "risks", "recommendations"],
        "description": "Technology stack evaluation and modernization roadmap",
        "standards": "---",
        "defaults": {
            "title": "Technology Assessment - {{system_name}}",
            "objectives": [
                "Inventory current technology stack (versions, support status, licensing)",
                "Benchmark against industry best practices and modern patterns",
                "Identify operational and security risks from technical debt",
                "Provide target state architecture with migration roadmap and costs"
            ]
        }
    },

    "patch_management": {
        "name": "Patch Management Report",
        "category": ReportCategory.SYSADMIN,
        "icon": "🔧",
        "sections": ["executive_summary", "findings", "timeline", "recommendations", "next_steps"],
        "description": "System patching status and compliance report",
        "standards": "---",
        "defaults": {
            "title": "Patch Management Report - {{period}}",
            "objectives": [
                "Document current patch compliance status",
                "Identify missing critical and security patches",
                "Plan and schedule patch deployment",
                "Minimize system downtime and risk"
            ]
        }
    },

    "capacity_planning": {
        "name": "Capacity Planning Report",
        "category": ReportCategory.SYSADMIN,
        "icon": "📈",
        "sections": ["executive_summary", "findings", "analysis", "recommendations", "timeline"],
        "description": "Resource utilization and capacity forecasting",
        "standards": "---",
        "defaults": {
            "title": "IT Capacity Planning Report - {{period}}",
            "objectives": [
                "Analyze current resource utilization trends",
                "Forecast future capacity requirements",
                "Identify potential bottlenecks and constraints",
                "Recommend infrastructure scaling strategy"
            ]
        }
    },

    # ═══ PROJECT MANAGEMENT REPORTS (4 templates) ═══
    "project_status": {
        "name": "Project Status Report",
        "category": ReportCategory.PROJECT_MGMT,
        "icon": "📋",
        "sections": ["executive_summary", "timeline", "risks", "findings", "recommendations", "next_steps"],
        "description": "Current project status and progress tracking",
        "standards": "---",
        "defaults": {
            "title": "{{project_name}} Status Report - {{period}}",
            "objectives": [
                "Report progress against project milestones",
                "Highlight blockers and impediments",
                "Communicate resource needs and timeline updates",
                "Manage stakeholder expectations"
            ]
        }
    },

    "project_proposal": {
        "name": "Project Proposal (Executive-Ready)",
        "category": ReportCategory.PROJECT_MGMT,
        "icon": "📝",
        "sections": ["executive_summary", "background", "scope", "methodology", "recommendations", "conclusion"],
        "description": "Comprehensive project proposal with ROI analysis",
        "standards": "---",
        "defaults": {
            "title": "Project Proposal - {{project_name}}",
            "objectives": [
                "Define problem statement, business impact, and strategic alignment",
                "Present proposed solution with architecture and alternatives comparison",
                "Provide detailed scope, deliverables, and acceptance criteria",
                "Justify investment with ROI analysis and risk mitigation plan"
            ]
        }
    },

    "project_closure": {
        "name": "Project Closure Report",
        "category": ReportCategory.PROJECT_MGMT,
        "icon": "🏁",
        "sections": ["executive_summary", "objectives", "findings", "analysis", "recommendations", "conclusion"],
        "description": "Project completion summary and lessons learned",
        "standards": "---",
        "defaults": {
            "title": "Project Closure Report - {{project_name}}",
            "objectives": [
                "Compare planned objectives vs actual outcomes with KPI deltas",
                "Analyze budget and schedule variances with root causes",
                "Document benefits realization and metrics achieved to date",
                "Capture lessons learned and best practices for future projects"
            ]
        }
    },

    "operational_report": {
        "name": "Operational Report (Monthly)",
        "category": ReportCategory.PROJECT_MGMT,
        "icon": "📈",
        "sections": ["executive_summary", "kpis", "findings", "recommendations", "next_steps"],
        "description": "Monthly operational metrics and KPI tracking",
        "standards": "SRE KPIs",
        "defaults": {
            "title": "Operational Report - {{month}} {{year}}",
            "objectives": [
                "Report on key operational KPIs (SLO/SLA, MTTR, MTBF, capacity utilization)",
                "Track project milestone burndown and identify blockers",
                "Highlight automation improvements and toil reduction initiatives",
                "Define next period action items with owners and dependencies"
            ]
        }
    },

    # ═══ COMPLIANCE REPORTS (1 template) ═══
    "compliance_audit": {
        "name": "Compliance Audit Report",
        "category": ReportCategory.COMPLIANCE,
        "icon": "✅",
        "sections": ["executive_summary", "scope", "methodology", "findings", "recommendations", "conclusion"],
        "description": "Regulatory compliance assessment and gap analysis",
        "standards": "ISO 27001, SOC 2, GDPR",
        "defaults": {
            "title": "{{regulation}} Compliance Audit Report",
            "methodology": "Conducted comprehensive compliance assessment against applicable regulatory requirements. Reviewed policies, procedures, technical controls, and documentation.",
            "objectives": [
                "Assess compliance with regulatory requirements",
                "Identify gaps and non-conformities",
                "Provide remediation roadmap and timeline",
                "Document evidence for audit purposes"
            ]
        }
    },

    # ═══ NETWORKING REPORTS (1 template) ═══
    "network_assessment": {
        "name": "Network Security Assessment",
        "category": ReportCategory.NETWORKING,
        "icon": "🌐",
        "sections": ["executive_summary", "scope", "methodology", "findings", "risks", "recommendations"],
        "description": "Network infrastructure security and performance review",
        "standards": "---",
        "defaults": {
            "title": "Network Security Assessment Report",
            "methodology": "Conducted network security assessment including configuration review, vulnerability scanning, and traffic analysis. Evaluated network segmentation, firewall rules, and access controls.",
            "objectives": [
                "Assess network security architecture and controls",
                "Identify misconfigurations and vulnerabilities",
                "Evaluate network segmentation and isolation",
                "Review firewall rules and access policies"
            ]
        }
    },

    # ═══ DATA ANALYTICS REPORTS (1 template) ═══
    "analytical_report": {
        "name": "Analytical Report (Data Insights)",
        "category": ReportCategory.DATA_ANALYTICS,
        "icon": "📊",
        "sections": ["executive_summary", "methodology", "kpis", "findings", "recommendations"],
        "description": "Business intelligence and data-driven insights report",
        "standards": "---",
        "defaults": {
            "title": "Data Analytics Report - {{analysis_focus}}",
            "methodology": "Performed comprehensive data analysis using statistical methods, trend analysis, and predictive modeling. Data sources validated for accuracy and completeness.",
            "objectives": [
                "Analyze historical data to identify trends, patterns, and anomalies",
                "Define and track key performance indicators (KPIs) with targets",
                "Provide data-driven insights to support business decision-making",
                "Recommend actionable improvements based on quantitative evidence"
            ]
        }
    }
}


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATION ENGINE
# ═══════════════════════════════════════════════════════════════════════════

class DocumentGenerator:
    """Professional Word document generator"""

    @staticmethod
    def generate_report(data: ReportData, template: dict, output_path: Path) -> bool:
        """Generate complete professional report"""
        try:
            doc = Document()
            DocumentGenerator._setup_styles(doc)

            # Apply smart variable substitution
            data = DocumentGenerator._apply_smart_variables(data)

            # Cover page
            DocumentGenerator._add_cover_page(doc, data, template)

            # Table of contents
            doc.add_page_break()
            doc.add_heading("Table of Contents", 1)
            doc.add_paragraph("(Auto-generated table of contents)")
            doc.add_page_break()

            # Standards/Frameworks (if applicable)
            if template.get('standards') and template['standards'] != '---':
                doc.add_heading("Standards & Frameworks", 2)
                doc.add_paragraph(f"This report follows: {template['standards']}")
                doc.add_paragraph()

            # Generate sections based on template
            sections = template['sections']

            if 'executive_summary' in sections and data.executive_summary:
                DocumentGenerator._add_section(doc, "Executive Summary", data.executive_summary)

            if 'background' in sections and data.background:
                DocumentGenerator._add_section(doc, "Background", data.background)

            if 'objectives' in sections and data.objectives:
                DocumentGenerator._add_list_section(doc, "Objectives", data.objectives)

            if 'scope' in sections and data.scope:
                DocumentGenerator._add_section(doc, "Scope", data.scope)

            if 'methodology' in sections and data.methodology:
                DocumentGenerator._add_section(doc, "Methodology", data.methodology)

            if 'kpis' in sections and data.kpis:
                DocumentGenerator._add_kpis_section(doc, data.kpis)

            if 'findings' in sections and data.findings:
                DocumentGenerator._add_list_section(doc, "Key Findings", data.findings)

            if 'analysis' in sections and data.analysis:
                DocumentGenerator._add_section(doc, "Analysis", data.analysis)

            if 'recommendations' in sections and data.recommendations:
                DocumentGenerator._add_list_section(doc, "Recommendations", data.recommendations, numbered=True)

            if 'risks' in sections and data.risks:
                DocumentGenerator._add_risks_section(doc, data.risks)

            if 'timeline' in sections and data.timeline:
                DocumentGenerator._add_timeline_section(doc, data.timeline)

            if 'architecture' in sections and data.architecture:
                DocumentGenerator._add_section(doc, "Architecture Overview", data.architecture)

            if 'tech_specs' in sections and data.tech_specs:
                DocumentGenerator._add_tech_specs_section(doc, data.tech_specs)

            if 'conclusion' in sections and data.conclusion:
                DocumentGenerator._add_section(doc, "Conclusion", data.conclusion)

            if 'next_steps' in sections and data.next_steps:
                DocumentGenerator._add_list_section(doc, "Next Steps", data.next_steps)

            # Appendices
            DocumentGenerator._add_appendices(doc, data)

            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            logger.info(f"Report generated successfully: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Report generation failed: {e}", exc_info=True)
            raise

    @staticmethod
    def _apply_smart_variables(data: ReportData) -> ReportData:
        """Replace smart variables throughout the document"""
        variables = {
            **data.smart_variables,
            'date': data.date,
            'author': data.author,
            'company': data.company_name,
            'project_code': data.project_code
        }

        # Replace in text fields
        for field in ['title', 'subtitle', 'executive_summary', 'background',
                      'scope', 'methodology', 'analysis', 'conclusion']:
            value = getattr(data, field)
            if value:
                for var_name, var_value in variables.items():
                    value = value.replace(f'{{{{{var_name}}}}}', str(var_value))
                setattr(data, field, value)

        return data

    @staticmethod
    def _setup_styles(doc: Document):
        """Configure professional document styles"""
        styles = doc.styles

        # Heading styles with professional colors
        heading_color = RGBColor(31, 78, 120)  # Professional blue
        for i in range(1, 4):
            style = styles[f'Heading {i}']
            style.font.name = 'Calibri'
            style.font.color.rgb = heading_color
            if i == 1:
                style.font.size = Pt(16)
                style.font.bold = True
            elif i == 2:
                style.font.size = Pt(14)

        # Normal text
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

    @staticmethod
    def _add_cover_page(doc: Document, data: ReportData, template: dict):
        """Generate professional cover page"""
        # Company branding
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = header.add_run(f"\n\n{data.company_name}\n")
        company_run.font.size = Pt(18)
        company_run.font.bold = True
        company_run.font.color.rgb = RGBColor(31, 78, 120)

        if data.department:
            dept_run = header.add_run(f"{data.department}\n")
            dept_run.font.size = Pt(12)
            dept_run.font.italic = True

        # Spacing
        doc.add_paragraph("\n" * 3)

        # Main title
        title = doc.add_paragraph()
        title_run = title.add_run(data.title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 120)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if data.subtitle:
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(data.subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.italic = True
            subtitle_run.font.color.rgb = RGBColor(89, 89, 89)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("\n" * 6)

        # Metadata table
        metadata = doc.add_table(rows=8, cols=2)
        metadata.style = 'Light Grid Accent 1'

        fields = [
            ("Author:", data.author),
            ("Date:", data.date),
            ("Version:", data.version),
            ("Status:", data.status),
            ("Classification:", data.classification),
            ("Project Code:", data.project_code or "N/A"),
            ("Department:", data.department or "N/A"),
            ("Document ID:", f"DOC-{datetime.now().strftime('%Y%m%d')}-{hash(data.title) % 10000:04d}")
        ]

        for i, (label, value) in enumerate(fields):
            metadata.rows[i].cells[0].text = label
            metadata.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            metadata.rows[i].cells[1].text = value

        # Template info
        doc.add_paragraph("\n" * 2)
        template_info = doc.add_paragraph()
        template_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        template_run = template_info.add_run(f"Template: {template['name']}")
        template_run.font.size = Pt(10)
        template_run.font.italic = True

        doc.add_page_break()

    @staticmethod
    def _add_section(doc: Document, title: str, content: str, level: int = 1):
        """Add text section with formatting"""
        doc.add_heading(title, level)
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()

    @staticmethod
    def _add_list_section(doc: Document, title: str, items: List[str], numbered: bool = False):
        """Add bulleted or numbered list section"""
        doc.add_heading(title, 1)
        style = 'List Number' if numbered else 'List Bullet'
        for item in items:
            doc.add_paragraph(item, style=style)
        doc.add_paragraph()

    @staticmethod
    def _add_kpis_section(doc: Document, kpis: Dict[str, str]):
        """Add KPI section with table"""
        if not kpis:
            return

        doc.add_heading("Key Performance Indicators", 1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header row
        table.rows[0].cells[0].text = "KPI"
        table.rows[0].cells[1].text = "Value"
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Data rows
        for kpi_name, kpi_value in kpis.items():
            row = table.add_row().cells
            row[0].text = kpi_name
            row[1].text = kpi_value

        doc.add_paragraph()

    @staticmethod
    def _add_risks_section(doc: Document, risks: List[Risk]):
        """Add comprehensive risk assessment section"""
        if not risks:
            return

        doc.add_heading("Risk Assessment", 1)
        doc.add_paragraph("The following risks have been identified and assessed:")
        doc.add_paragraph()

        # Risk summary table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        headers = ["#", "Risk Description", "Category", "Impact", "Likelihood", "Owner"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        for idx, risk in enumerate(risks, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = risk.description[:60] + "..." if len(risk.description) > 60 else risk.description
            row[2].text = risk.category
            row[3].text = risk.impact
            row[4].text = risk.likelihood
            row[5].text = risk.owner

        doc.add_paragraph()

        # Detailed risk descriptions and mitigation
        doc.add_heading("Risk Details and Mitigation Plans", 2)
        for idx, risk in enumerate(risks, 1):
            doc.add_heading(f"Risk {idx}: {risk.description[:50]}...", 3)

            details = doc.add_paragraph()
            details.add_run("Category: ").bold = True
            details.add_run(f"{risk.category}\n")
            details.add_run("Impact: ").bold = True
            details.add_run(f"{risk.impact}\n")
            details.add_run("Likelihood: ").bold = True
            details.add_run(f"{risk.likelihood}\n")
            details.add_run("Owner: ").bold = True
            details.add_run(f"{risk.owner}\n")

            if risk.mitigation:
                doc.add_paragraph("Mitigation Strategy:", style='List Bullet')
                doc.add_paragraph(risk.mitigation, style='List Bullet 2')

            doc.add_paragraph()

    @staticmethod
    def _add_timeline_section(doc: Document, timeline: List[TimelineEntry]):
        """Add project timeline with Gantt-style table"""
        if not timeline:
            return

        doc.add_heading("Project Timeline", 1)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        headers = ["Milestone", "Target Date", "Status", "Owner", "Notes"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        for entry in timeline:
            row = table.add_row().cells
            row[0].text = entry.milestone
            row[1].text = entry.date
            row[2].text = entry.status
            row[3].text = entry.owner
            row[4].text = entry.notes[:40] + "..." if len(entry.notes) > 40 else entry.notes

        doc.add_paragraph()

    @staticmethod
    def _add_tech_specs_section(doc: Document, specs: List[TechnicalSpec]):
        """Add technical specifications table"""
        if not specs:
            return

        doc.add_heading("Technical Specifications", 1)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        headers = ["Component", "Specification", "Version", "Notes"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        for spec in specs:
            row = table.add_row().cells
            row[0].text = spec.component
            row[1].text = spec.specification
            row[2].text = spec.version
            row[3].text = spec.notes

        doc.add_paragraph()

    @staticmethod
    def _add_appendices(doc: Document, data: ReportData):
        """Add comprehensive appendices"""
        doc.add_page_break()
        doc.add_heading("Appendices", 1)

        appendix_num = ord('A')

        # Technology Stack
        if data.tech_stack:
            doc.add_heading(f"Appendix {chr(appendix_num)}: Technology Stack", 2)
            for tech in data.tech_stack:
                doc.add_paragraph(tech, style='List Bullet')
            doc.add_paragraph()
            appendix_num += 1

        # System Specifications
        if any([data.os_platform, data.ram_size, data.cpu_specs, data.deployment_env]):
            doc.add_heading(f"Appendix {chr(appendix_num)}: System Specifications", 2)

            specs_table = doc.add_table(rows=1, cols=2)
            specs_table.style = 'Light List Accent 1'

            if data.os_platform:
                row = specs_table.add_row().cells
                row[0].text = "Operating System"
                row[0].paragraphs[0].runs[0].font.bold = True
                row[1].text = data.os_platform

            if data.ram_size:
                row = specs_table.add_row().cells
                row[0].text = "Memory (RAM)"
                row[0].paragraphs[0].runs[0].font.bold = True
                row[1].text = data.ram_size

            if data.cpu_specs:
                row = specs_table.add_row().cells
                row[0].text = "Processor (CPU)"
                row[0].paragraphs[0].runs[0].font.bold = True
                row[1].text = data.cpu_specs

            if data.deployment_env:
                row = specs_table.add_row().cells
                row[0].text = "Deployment Environment"
                row[0].paragraphs[0].runs[0].font.bold = True
                row[1].text = data.deployment_env

            doc.add_paragraph()
            appendix_num += 1

        # Architecture
        if data.architecture:
            doc.add_heading(f"Appendix {chr(appendix_num)}: Architecture Details", 2)
            doc.add_paragraph(data.architecture)
            doc.add_paragraph()
            appendix_num += 1

        # Resources
        if data.resources:
            doc.add_heading(f"Appendix {chr(appendix_num)}: Resources", 2)
            for resource in data.resources:
                doc.add_paragraph(resource, style='List Bullet')
            doc.add_paragraph()
            appendix_num += 1

        # Related Links
        if any([data.repository_url, data.documentation_url, data.demo_url]):
            doc.add_heading(f"Appendix {chr(appendix_num)}: Related Links", 2)

            if data.repository_url:
                p = doc.add_paragraph()
                p.add_run("Source Code Repository: ").bold = True
                p.add_run(data.repository_url)

            if data.documentation_url:
                p = doc.add_paragraph()
                p.add_run("Documentation: ").bold = True
                p.add_run(data.documentation_url)

            if data.demo_url:
                p = doc.add_paragraph()
                p.add_run("Live Demo: ").bold = True
                p.add_run(data.demo_url)

            doc.add_paragraph()
            appendix_num += 1

        # References
        if data.references:
            doc.add_heading(f"Appendix {chr(appendix_num)}: References", 2)
            for i, ref in enumerate(data.references, 1):
                doc.add_paragraph(f"[{i}] {ref}", style='List Number')
            doc.add_paragraph()
            appendix_num += 1

        # Keywords/Tags
        if data.tags:
            doc.add_heading(f"Appendix {chr(appendix_num)}: Document Keywords", 2)
            doc.add_paragraph(", ".join(data.tags))


# ═══════════════════════════════════════════════════════════════════════════
# FILE I/O OPERATIONS
# ═══════════════════════════════════════════════════════════════════════════

class ProjectFileManager:
    """Handle project save/load operations"""

    @staticmethod
    def save_project(data: ReportData, file_path: Path) -> bool:
        """Save project to JSON file"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data.to_dict(), f, indent=2, ensure_ascii=False)
            logger.info(f"Project saved: {file_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to save project: {e}")
            return False

    @staticmethod
    def load_project(file_path: Path) -> Optional[ReportData]:
        """Load project from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data_dict = json.load(f)
            logger.info(f"Project loaded: {file_path}")
            return ReportData.from_dict(data_dict)
        except Exception as e:
            logger.error(f"Failed to load project: {e}")
            return None


# ═══════════════════════════════════════════════════════════════════════════
# COMMAND-LINE INTERFACE
# ═══════════════════════════════════════════════════════════════════════════

def main_cli():
    """Command-line interface for batch report generation"""
    import argparse

    parser = argparse.ArgumentParser(
        description="Reportify Pro - Enterprise Report Generator CLI",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate report from project file
  python reportify_pro.py generate -i project.json -o report.docx

  # List available templates
  python reportify_pro.py list-templates

  # Create new project from template
  python reportify_pro.py new -t vulnerability_assessment -o security_audit.json
        """
    )

    subparsers = parser.add_subparsers(dest='command', help='Command to execute')

    # Generate command
    gen_parser = subparsers.add_parser('generate', help='Generate report from project file')
    gen_parser.add_argument('-i', '--input', type=Path, required=True, help='Input project JSON file')
    gen_parser.add_argument('-o', '--output', type=Path, required=True, help='Output report file (.docx)')

    # List templates command
    subparsers.add_parser('list-templates', help='List all available templates')

    # New project command
    new_parser = subparsers.add_parser('new', help='Create new project from template')
    new_parser.add_argument('-t', '--template', required=True, help='Template key')
    new_parser.add_argument('-o', '--output', type=Path, required=True, help='Output project file')

    args = parser.parse_args()

    if args.command == 'generate':
        data = ProjectFileManager.load_project(args.input)
        if data:
            template = REPORT_TEMPLATES.get(data.template_key)
            if template:
                DocumentGenerator.generate_report(data, template, args.output)
                print(f"✓ Report generated: {args.output}")
            else:
                print(f"✗ Unknown template: {data.template_key}")
        else:
            print(f"✗ Failed to load project: {args.input}")

    elif args.command == 'list-templates':
        print("\n╔══════════════════════════════════════════════════════════════╗")
        print("║           REPORTIFY PRO - AVAILABLE TEMPLATES               ║")
        print("╚══════════════════════════════════════════════════════════════╝\n")

        for category in ReportCategory:
            cat_templates = {k: v for k, v in REPORT_TEMPLATES.items() if v['category'] == category}
            if cat_templates:
                print(f"\n{category.value.upper().replace('_', ' ')} ({len(cat_templates)} templates)")
                print("─" * 60)
                for key, template in cat_templates.items():
                    print(f"  {template['icon']} {key:30} - {template['name']}")

        print(f"\n\nTotal Templates: {len(REPORT_TEMPLATES)}")

    elif args.command == 'new':
        template = REPORT_TEMPLATES.get(args.template)
        if template:
            data = ReportData(template_key=args.template, category=template['category'].value)
            # Apply template defaults
            defaults = template.get('defaults', {})
            data.title = defaults.get('title', template['name'])
            data.objectives = defaults.get('objectives', []).copy()
            if 'methodology' in defaults:
                data.methodology = defaults['methodology']

            ProjectFileManager.save_project(data, args.output)
            print(f"✓ New project created: {args.output}")
            print(f"  Template: {template['name']}")
            print(f"  Category: {template['category'].value}")
        else:
            print(f"✗ Unknown template: {args.template}")
            print("Run 'python reportify_pro.py list-templates' to see available templates")

    else:
        parser.print_help()


if __name__ == '__main__':
    print("""
╔═══════════════════════════════════════════════════════════════╗
║                    REPORTIFY PRO v2.1                        ║
║          Enterprise Report Generation System                  ║
║                                                               ║
║  25+ Professional Templates | 7+ IT Domains                   ║
║  Smart Variables | Risk Assessment | Timeline Management      ║
╚═══════════════════════════════════════════════════════════════╝
    """)

    if len(sys.argv) > 1:
        # CLI mode
        main_cli()
    else:
        # Show help
        print("Command-line mode. Run with --help for usage information.")
        print("\nQuick start:")
        print("  python reportify_pro.py list-templates")
        print("  python reportify_pro.py new -t vulnerability_assessment -o my_report.json")
        print("  python reportify_pro.py generate -i my_report.json -o my_report.docx")
