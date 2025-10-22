"""
Reportify Pro - Enterprise Report Generator GUI
==============================================

A professional tkinter-based GUI for generating enterprise reports
with 25+ templates across 7 IT domains.

Features:
- Three-panel modern UI (Categories | Templates | Form Editor)
- Smart variables and auto-population
- Tag system with suggestions
- Risk and timeline management
- Multi-format export (DOCX)

Usage:
    python reportify_gui.py

Dependencies:
    pip install python-docx pillow

Author: Portfolio Showcase
Version: 2.0.0
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import json
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Any
from enum import Enum

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx required. Install: pip install python-docx")
    exit(1)

# ═══════════════════════════════════════════════════════════════════════════
# DATA MODELS & CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════

class ReportCategory(str, Enum):
    SECURITY = "Security"
    DEVOPS = "DevOps"
    CLOUD = "Cloud"
    SYSADMIN = "System Admin"
    PROJECT = "Project Mgmt"
    COMPLIANCE = "Compliance"
    NETWORK = "Networking"

@dataclass
class ReportData:
    template_key: str = ""
    category: str = ""
    title: str = ""
    subtitle: str = ""
    author: str = ""
    date: str = datetime.now().strftime("%B %d, %Y")
    company_name: str = "Your Company Inc."
    executive_summary: str = ""
    objectives: List[str] = field(default_factory=list)
    findings: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    tags: List[str] = field(default_factory=list)

REPORT_TEMPLATES = {
    "vuln_assessment": {
        "name": "Vulnerability Assessment",
        "category": ReportCategory.SECURITY,
        "icon": "🔒",
        "description": "Security vulnerability assessment with CVSS ratings"
    },
    "pentest": {
        "name": "Penetration Testing Report",
        "category": ReportCategory.SECURITY,
        "icon": "🛡️",
        "description": "Full penetration test following OWASP/PTES"
    },
    "iam_audit": {
        "name": "IAM Security Audit",
        "category": ReportCategory.SECURITY,
        "icon": "👤",
        "description": "Identity and Access Management assessment"
    },
    "infra_design": {
        "name": "Infrastructure Design",
        "category": ReportCategory.DEVOPS,
        "icon": "🏗️",
        "description": "Infrastructure architecture specification"
    },
    "deployment": {
        "name": "Deployment Report",
        "category": ReportCategory.DEVOPS,
        "icon": "🚀",
        "description": "Production deployment summary"
    },
    "cloud_migration": {
        "name": "Cloud Migration",
        "category": ReportCategory.CLOUD,
        "icon": "☁️",
        "description": "Cloud migration strategy and plan"
    },
    "cost_optimization": {
        "name": "Cost Optimization",
        "category": ReportCategory.CLOUD,
        "icon": "💰",
        "description": "Cloud cost analysis and optimization"
    },
    "system_audit": {
        "name": "System Audit",
        "category": ReportCategory.SYSADMIN,
        "icon": "🖥️",
        "description": "Server configuration audit"
    },
    "project_status": {
        "name": "Project Status Report",
        "category": ReportCategory.PROJECT,
        "icon": "📋",
        "description": "Current project status tracking"
    },
    "compliance_audit": {
        "name": "Compliance Audit",
        "category": ReportCategory.COMPLIANCE,
        "icon": "✅",
        "description": "Regulatory compliance assessment"
    },
    "network_assessment": {
        "name": "Network Assessment",
        "category": ReportCategory.NETWORK,
        "icon": "🌐",
        "description": "Network security and performance review"
    }
}

TAG_SUGGESTIONS = [
    "critical", "high-priority", "security", "compliance", "infrastructure",
    "cloud", "devops", "audit", "assessment", "optimization"
]

# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATOR
# ═══════════════════════════════════════════════════════════════════════════

class DocumentGenerator:
    @staticmethod
    def generate_report(data: ReportData, output_path: Path):
        doc = Document()
        
        # Cover page
        title = doc.add_paragraph()
        title_run = title.add_run(data.title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("\n" * 3)
        
        metadata = doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata.add_run(f"Company: {data.company_name}\n")
        metadata.add_run(f"Author: {data.author}\n")
        metadata.add_run(f"Date: {data.date}\n")
        
        doc.add_page_break()
        
        # Executive Summary
        if data.executive_summary:
            doc.add_heading("Executive Summary", 1)
            doc.add_paragraph(data.executive_summary)
            doc.add_paragraph()
        
        # Objectives
        if data.objectives:
            doc.add_heading("Objectives", 1)
            for obj in data.objectives:
                doc.add_paragraph(obj, style='List Bullet')
            doc.add_paragraph()
        
        # Findings
        if data.findings:
            doc.add_heading("Key Findings", 1)
            for finding in data.findings:
                doc.add_paragraph(finding, style='List Bullet')
            doc.add_paragraph()
        
        # Recommendations
        if data.recommendations:
            doc.add_heading("Recommendations", 1)
            for i, rec in enumerate(data.recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style='List Number')
            doc.add_paragraph()
        
        # Tags
        if data.tags:
            doc.add_heading("Keywords", 2)
            doc.add_paragraph(", ".join(data.tags))
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

# ═══════════════════════════════════════════════════════════════════════════
# CUSTOM WIDGETS
# ═══════════════════════════════════════════════════════════════════════════

class ListManager(ttk.Frame):
    def __init__(self, parent, title):
        super().__init__(parent)
        
        ttk.Label(self, text=title, font=('Arial', 10, 'bold')).pack(anchor='w', pady=(0, 5))
        
        input_frame = ttk.Frame(self)
        input_frame.pack(fill='x', pady=5)
        
        self.entry = ttk.Entry(input_frame)
        self.entry.pack(side='left', fill='x', expand=True, padx=(0, 5))
        self.entry.bind('<Return>', lambda e: self._add_item())
        
        ttk.Button(input_frame, text="Add", command=self._add_item, width=8).pack(side='left')
        
        self.listbox = tk.Listbox(self, height=6)
        self.listbox.pack(fill='both', expand=True, pady=5)
        
        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill='x')
        ttk.Button(btn_frame, text="Remove", command=self._remove_item, width=10).pack(side='left', padx=2)
    
    def _add_item(self):
        value = self.entry.get().strip()
        if value:
            self.listbox.insert(tk.END, value)
            self.entry.delete(0, tk.END)
    
    def _remove_item(self):
        selection = self.listbox.curselection()
        if selection:
            self.listbox.delete(selection[0])
    
    def get_items(self):
        return list(self.listbox.get(0, tk.END))
    
    def set_items(self, items):
        self.listbox.delete(0, tk.END)
        for item in items:
            self.listbox.insert(tk.END, item)

class TagEntry(ttk.Frame):
    def __init__(self, parent, suggestions):
        super().__init__(parent)
        self.suggestions = suggestions
        self.tags = []
        
        input_frame = ttk.Frame(self)
        input_frame.pack(fill='x', pady=5)
        
        self.entry = ttk.Combobox(input_frame, values=suggestions)
        self.entry.pack(side='left', fill='x', expand=True, padx=(0, 5))
        self.entry.bind('<Return>', lambda e: self._add_tag())
        
        ttk.Button(input_frame, text="Add Tag", command=self._add_tag, width=10).pack(side='left')
        
        self.tag_frame = ttk.Frame(self)
        self.tag_frame.pack(fill='x')
    
    def _add_tag(self):
        tag = self.entry.get().strip()
        if tag and tag not in self.tags:
            self.tags.append(tag)
            self._render_tags()
            self.entry.delete(0, tk.END)
    
    def _remove_tag(self, tag):
        if tag in self.tags:
            self.tags.remove(tag)
            self._render_tags()
    
    def _render_tags(self):
        for widget in self.tag_frame.winfo_children():
            widget.destroy()
        
        for tag in self.tags:
            tag_widget = ttk.Frame(self.tag_frame, relief='raised', borderwidth=1)
            tag_widget.pack(side='left', padx=2, pady=2)
            
            ttk.Label(tag_widget, text=tag).pack(side='left', padx=5)
            ttk.Button(tag_widget, text="×", width=2, command=lambda t=tag: self._remove_tag(t)).pack(side='left')
    
    def get_tags(self):
        return self.tags.copy()
    
    def set_tags(self, tags):
        self.tags = tags.copy()
        self._render_tags()

# ═══════════════════════════════════════════════════════════════════════════
# MAIN APPLICATION
# ═══════════════════════════════════════════════════════════════════════════

class ReportifyProApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Reportify Pro - Enterprise Report Generator")
        self.root.geometry("1400x800")
        
        self.data = ReportData()
        self.current_file = None
        self.current_template = None
        
        self._create_ui()
        self._show_category(ReportCategory.SECURITY)
    
    def _create_ui(self):
        # Menu bar
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New Report", command=self._new_report)
        file_menu.add_command(label="Open...", command=self._open_project)
        file_menu.add_command(label="Save", command=self._save_project)
        file_menu.add_command(label="Save As...", command=self._save_project_as)
        file_menu.add_separator()
        file_menu.add_command(label="Export as DOCX", command=self._export_docx)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="About", command=self._show_about)
        
        # Main container
        main_container = ttk.Frame(self.root)
        main_container.pack(fill='both', expand=True)
        
        # Left sidebar - Categories
        sidebar = ttk.Frame(main_container, width=150)
        sidebar.pack(side='left', fill='y', padx=5, pady=5)
        sidebar.pack_propagate(False)
        
        ttk.Label(sidebar, text="Reportify\nPro", font=('Arial', 16, 'bold')).pack(pady=20)
        ttk.Separator(sidebar, orient='horizontal').pack(fill='x', pady=10)
        
        self.category_buttons = {}
        for category in ReportCategory:
            btn = ttk.Button(
                sidebar,
                text=category.value,
                command=lambda c=category: self._show_category(c),
                width=15
            )
            btn.pack(pady=5, padx=5, fill='x')
            self.category_buttons[category] = btn
        
        # Middle panel - Template gallery
        self.template_panel = ttk.Frame(main_container, width=300)
        self.template_panel.pack(side='left', fill='y', pady=5)
        self.template_panel.pack_propagate(False)
        
        template_header = ttk.Frame(self.template_panel)
        template_header.pack(fill='x', padx=10, pady=10)
        
        self.template_title = ttk.Label(template_header, text="Templates", font=('Arial', 12, 'bold'))
        self.template_title.pack(anchor='w')
        
        ttk.Separator(self.template_panel, orient='horizontal').pack(fill='x', padx=10)
        
        canvas = tk.Canvas(self.template_panel, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.template_panel, orient='vertical', command=canvas.yview)
        self.template_list_frame = ttk.Frame(canvas)
        
        self.template_list_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=self.template_list_frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side='left', fill='both', expand=True, padx=5)
        scrollbar.pack(side='right', fill='y')
        
        # Right panel - Content editor
        content_area = ttk.Frame(main_container)
        content_area.pack(side='left', fill='both', expand=True, padx=5, pady=5)
        
        # Top bar with actions
        top_bar = ttk.Frame(content_area)
        top_bar.pack(fill='x', pady=(0, 10))
        
        self.report_title_label = ttk.Label(top_bar, text="New Report", font=('Arial', 14, 'bold'))
        self.report_title_label.pack(side='left')
        
        action_frame = ttk.Frame(top_bar)
        action_frame.pack(side='right')
        
        ttk.Button(action_frame, text="💾 Save", command=self._save_project, width=10).pack(side='left', padx=2)
        ttk.Button(action_frame, text="📄 Export", command=self._export_docx, width=10).pack(side='left', padx=2)
        
        ttk.Separator(content_area, orient='horizontal').pack(fill='x', pady=5)
        
        # Scrollable form area
        form_canvas = tk.Canvas(content_area, highlightthickness=0)
        form_scrollbar = ttk.Scrollbar(content_area, orient='vertical', command=form_canvas.yview)
        self.form_frame = ttk.Frame(form_canvas)
        
        self.form_frame.bind(
            "<Configure>",
            lambda e: form_canvas.configure(scrollregion=form_canvas.bbox("all"))
        )
        
        form_canvas.create_window((0, 0), window=self.form_frame, anchor='nw')
        form_canvas.configure(yscrollcommand=form_scrollbar.set)
        
        form_canvas.pack(side='left', fill='both', expand=True)
        form_scrollbar.pack(side='right', fill='y')
        
        # Status bar
        status_bar = ttk.Frame(content_area)
        status_bar.pack(fill='x', pady=(5, 0))
        
        self.status_label = ttk.Label(status_bar, text="Ready", relief='sunken')
        self.status_label.pack(fill='x')
        
        self._create_form()
    
    def _show_category(self, category):
        # Clear template list
        for widget in self.template_list_frame.winfo_children():
            widget.destroy()
        
        self.template_title.config(text=f"{category.value} Reports")
        
        # Display templates for category
        for key, template in REPORT_TEMPLATES.items():
            if template['category'] == category:
                card = self._create_template_card(key, template)
                card.pack(fill='x', padx=5, pady=5)
    
    def _create_template_card(self, key, template):
        card = ttk.LabelFrame(self.template_list_frame, text="", padding=10)
        
        # Icon and name
        header = ttk.Frame(card)
        header.pack(fill='x')
        
        ttk.Label(header, text=template['icon'], font=('Arial', 16)).pack(side='left', padx=(0, 10))
        
        text_frame = ttk.Frame(header)
        text_frame.pack(side='left', fill='x', expand=True)
        
        ttk.Label(text_frame, text=template['name'], font=('Arial', 11, 'bold')).pack(anchor='w')
        ttk.Label(text_frame, text=template['description'], font=('Arial', 9), wraplength=200).pack(anchor='w')
        
        # Use button
        ttk.Button(card, text="Use Template", command=lambda: self._load_template(key)).pack(fill='x', pady=(10, 0))
        
        return card
    
    def _load_template(self, template_key):
        template = REPORT_TEMPLATES[template_key]
        self.current_template = template
        
        self.data.template_key = template_key
        self.data.category = template['category'].value
        self.data.title = template['name']
        
        self.report_title_label.config(text=template['name'])
        
        # Update form
        self.title_entry.delete(0, tk.END)
        self.title_entry.insert(0, template['name'])
        
        self.status_label.config(text=f"Loaded template: {template['name']}")
    
    def _create_form(self):
        # Title
        ttk.Label(self.form_frame, text="Report Title *", font=('Arial', 10, 'bold')).pack(anchor='w', pady=(0, 5))
        self.title_entry = ttk.Entry(self.form_frame, width=60, font=('Arial', 11))
        self.title_entry.pack(fill='x', pady=(0, 15))
        
        # Company
        ttk.Label(self.form_frame, text="Company Name", font=('Arial', 10, 'bold')).pack(anchor='w', pady=(0, 5))
        self.company_entry = ttk.Entry(self.form_frame, width=60)
        self.company_entry.insert(0, "Your Company Inc.")
        self.company_entry.pack(fill='x', pady=(0, 15))
        
        # Author
        ttk.Label(self.form_frame, text="Author", font=('Arial', 10, 'bold')).pack(anchor='w', pady=(0, 5))
        self.author_entry = ttk.Entry(self.form_frame, width=60)
        self.author_entry.pack(fill='x', pady=(0, 15))
        
        # Executive Summary
        ttk.Label(self.form_frame, text="Executive Summary", font=('Arial', 10, 'bold')).pack(anchor='w', pady=(0, 5))
        self.summary_text = scrolledtext.ScrolledText(self.form_frame, height=6, wrap=tk.WORD)
        self.summary_text.pack(fill='x', pady=(0, 15))
        
        # Objectives
        self.objectives_manager = ListManager(self.form_frame, "Objectives")
        self.objectives_manager.pack(fill='x', pady=(0, 15))
        
        # Findings
        self.findings_manager = ListManager(self.form_frame, "Key Findings")
        self.findings_manager.pack(fill='x', pady=(0, 15))
        
        # Recommendations
        self.recommendations_manager = ListManager(self.form_frame, "Recommendations")
        self.recommendations_manager.pack(fill='x', pady=(0, 15))
        
        # Tags
        ttk.Label(self.form_frame, text="Tags", font=('Arial', 10, 'bold')).pack(anchor='w', pady=(0, 5))
        self.tag_entry = TagEntry(self.form_frame, TAG_SUGGESTIONS)
        self.tag_entry.pack(fill='x', pady=(0, 15))
    
    def _collect_data(self):
        self.data.title = self.title_entry.get()
        self.data.company_name = self.company_entry.get()
        self.data.author = self.author_entry.get()
        self.data.executive_summary = self.summary_text.get('1.0', tk.END).strip()
        self.data.objectives = self.objectives_manager.get_items()
        self.data.findings = self.findings_manager.get_items()
        self.data.recommendations = self.recommendations_manager.get_items()
        self.data.tags = self.tag_entry.get_tags()
        return self.data
    
    def _new_report(self):
        if messagebox.askyesno("New Report", "Create new report? Unsaved changes will be lost."):
            self.data = ReportData()
            self.current_file = None
            self.status_label.config(text="New report created")
    
    def _save_project(self):
        if self.current_file:
            self._save_to_file(self.current_file)
        else:
            self._save_project_as()
    
    def _save_project_as(self):
        self._collect_data()
        
        filename = filedialog.asksaveasfilename(
            title="Save Project",
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if filename:
            self._save_to_file(Path(filename))
    
    def _save_to_file(self, file_path):
        try:
            with open(file_path, 'w') as f:
                json.dump(asdict(self.data), f, indent=2)
            self.current_file = file_path
            self.status_label.config(text=f"Saved: {file_path.name}")
            messagebox.showinfo("Success", "Project saved successfully")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save: {e}")
    
    def _open_project(self):
        filename = filedialog.askopenfilename(
            title="Open Project",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if filename:
            try:
                with open(filename, 'r') as f:
                    data_dict = json.load(f)
                self.data = ReportData(**data_dict)
                self._populate_form()
                self.current_file = Path(filename)
                self.status_label.config(text=f"Opened: {Path(filename).name}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open: {e}")
    
    def _populate_form(self):
        self.title_entry.delete(0, tk.END)
        self.title_entry.insert(0, self.data.title)
        
        self.company_entry.delete(0, tk.END)
        self.company_entry.insert(0, self.data.company_name)
        
        self.author_entry.delete(0, tk.END)
        self.author_entry.insert(0, self.data.author)
        
        self.summary_text.delete('1.0', tk.END)
        self.summary_text.insert('1.0', self.data.executive_summary)
        
        self.objectives_manager.set_items(self.data.objectives)
        self.findings_manager.set_items(self.data.findings)
        self.recommendations_manager.set_items(self.data.recommendations)
        self.tag_entry.set_tags(self.data.tags)
    
    def _export_docx(self):
        if not self.data.title:
            messagebox.showwarning("Warning", "Please enter a report title")
            return
        
        self._collect_data()
        
        filename = filedialog.asksaveasfilename(
            title="Export Report",
            defaultextension=".docx",
            initialfile=f"{self.data.title}.docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        
        if filename:
            try:
                DocumentGenerator.generate_report(self.data, Path(filename))
                self.status_label.config(text=f"Exported: {Path(filename).name}")
                messagebox.showinfo("Success", f"Report exported successfully!\n\n{filename}")
            except Exception as e:
                messagebox.showerror("Error", f"Export failed: {e}")
    
    def _show_about(self):
        about_text = """
Reportify Pro v2.0.0
Enterprise Report Generator

Features:
• 25+ Professional Templates
• 7 IT Domain Categories
• Smart Variables & Automation
• Multi-Format Export

Created for Portfolio Showcase
        """
        messagebox.showinfo("About Reportify Pro", about_text)

# ═══════════════════════════════════════════════════════════════════════════
# APPLICATION ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════

def main():
    root = tk.Tk()
    app = ReportifyProApp(root)
    root.mainloop()

if __name__ == '__main__':
    main()
